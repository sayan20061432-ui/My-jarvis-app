import os
import requests
from kivy.lang import Builder
from kivymd.app import MDApp
from google import genai
from google.genai import types
from jnius import autoclass
from reportlab.pdfgen import canvas
from openpyxl import Workbook
from docx import Document

# ==========================================
# 🔑 INSERT YOUR CONFIGURATION HERE 🔑
# ==========================================
GEMINI_KEY = "AQ.Ab8RN6L8e7MmN8C3d6kQSmudhgQlp4u4ZqbEPaCtG5Y-UiWE3A"
ELEVEN_KEY = "sk_4538afbcf86db0ac191415c864cc2bbe809e3f94a0cdbf1c"
VOICE_ID = "ZA3TGoYAsdYMffXndkSX"
# ==========================================

KV = '''
MDScreen:
    md_bg_color: 0.01, 0.01, 0.04, 1
    MDBoxLayout:
        orientation: 'vertical'
        padding: 25
        spacing: 20
        MDLabel:
            text: "J.A.R.V.I.S. OS V1.0"
            halign: "center"
            font_style: "H5"
            theme_text_color: "Custom"
            text_color: 0, 0.85, 1, 1
        MDLabel:
            id: console_output
            text: "All core systems initialized, Sir. Awaiting your instruction."
            halign: "left"
            font_style: "Body1"
            theme_text_color: "Custom"
            text_color: 0.75, 0.9, 1, 1
            size_hint_y: 0.6
            valign: "top"
        MDTextField:
            id: user_input
            hint_text: "Input command..."
            line_color_focus: 0, 0.85, 1, 1
            text_color_focus: 1, 1, 1, 1
            hint_text_color_focus: 0, 0.85, 1, 1
        MDFloatingActionButton:
            icon: "send"
            md_bg_color: 0, 0.6, 0.9, 1
            icon_color: 1, 1, 1, 1
            pos_hint: {"center_x": .5}
            on_release: app.process_command()
'''

class JarvisApp(MDApp):
    def build(self):
        self.theme_cls.theme_style = "Dark"
        self.ai_client = genai.Client(api_key=GEMINI_KEY)
        return Builder.load_string(KV)
        
    def process_command(self):
        query = self.root.ids.user_input.text.strip()
        if not query:
            return
            
        self.root.ids.console_output.text = "Processing request, Sir..."
        response_text = ""
        
        if query.lower().startswith("open "):
            app_name = query.lower().replace("open ", "").strip()
            response_text = self.launch_android_app(app_name)
            
        elif "create pdf file" in query.lower():
            response_text = self.make_pdf_document()
        elif "create word file" in query.lower():
            response_text = self.make_word_document()
        elif "create excel file" in query.lower():
            response_text = self.make_excel_document()
            
        else:
            system_prompt = (
                "You are JARVIS, the legendary AI assistant from Iron Man. "
                "Always address the user as 'Sir'. Keep your responses sharp, sophisticated, and concise."
            )
            try:
                response = self.ai_client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=query,
                    config=types.GenerateContentConfig(
                        system_instruction=system_prompt,
                        tools=[{"google_search": {}}],
                    )
                )
                response_text = response.text
            except Exception as e:
                response_text = f"Apologies Sir, I encountered an external neural link error: {str(e)}"
                
        self.root.ids.console_output.text = f"JARVIS:\n{response_text}"
        self.root.ids.user_input.text = ""
        self.speak_via_elevenlabs(response_text)

    def launch_android_app(self, app_name):
        try:
            PythonActivity = autoclass('org.kivy.android.PythonActivity')
            pm = PythonActivity.mActivity.getPackageManager()
            packages = pm.getInstalledPackages(0)
            for i in range(packages.size()):
                pkg_info = packages.get(i)
                label = str(pkg_info.applicationInfo.loadLabel(pm))
                if app_name in label.lower():
                    intent = pm.getLaunchIntentForPackage(pkg_info.packageName)
                    if intent:
                        PythonActivity.mActivity.startActivity(intent)
                        return f"Accessing the target network matrix. Opening {label} now, Sir."
            return f"I am unable to locate an application named '{app_name}' within the local mainframe, Sir."
        except Exception:
            return f"Emulation limitation. Application launch system requires deployment on the mobile device, Sir."

    def make_pdf_document(self):
        try:
            filename = "/sdcard/Download/Jarvis_Briefing.pdf"
            c = canvas.Canvas(filename)
            c.drawString(100, 750, "J.A.R.V.I.S. SECURE BRIEFING")
            c.drawString(100, 700, "Systems operational. Database architecture integrity verified at 100%.")
            c.save()
            return f"Secure data compilation complete. PDF document generated at {filename}, Sir."
        except Exception as e:
            return f"Failed to finalize file matrix compilation: {str(e)}"

    def make_word_document(self):
        try:
            filename = "/sdcard/Download/Jarvis_Report.docx"
            doc = Document()
            doc.add_heading('J.A.R.V.I.S. Core Log', level=1)
            doc.add_paragraph('Automated report generated by your mobile assistant mainframe.')
            doc.save(filename)
            return f"Word document compiled and exported to {filename}, Sir."
        except Exception as e:
            return f"Error building structural text file: {str(e)}"

    def make_excel_document(self):
        try:
            filename = "/sdcard/Download/Jarvis_Data.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.append(["Timestamp", "System Component", "Status"])
            ws.append(["00:00", "Neural Net", "Optimal"])
            wb.save(filename)
            return f"Spreadsheet metrics calculation complete. Exported to {filename}, Sir."
        except Exception as e:
            return f"Error building structural spreadsheet file: {str(e)}"

    def speak_via_elevenlabs(self, text):
        clean_text = text.replace("*", "").replace("#", "")
        url = f"https://elevenlabs.io{VOICE_ID}/stream"
        headers = {"xi-api-key": ELEVEN_KEY, "Content-Type": "application/json"}
        data = {
            "text": clean_text,
            "model_id": "eleven_monolingual_v1",
            "voice_settings": {"stability": 0.75, "similarity_boost": 0.85}
        }
        try:
            response = requests.post(url, json=data, headers=headers, stream=True)
            if response.status_code == 200:
                with open("/sdcard/Download/jarvis_voice.mp3", "wb") as f:
                    for chunk in response.iter_content(chunk_size=1024):
                        if chunk: f.write(chunk)
        except Exception:
            pass

if __name__ == '__main__':
    JarvisApp().run()
